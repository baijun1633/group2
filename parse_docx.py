import docx
import re

def parse_docx_to_markdown(docx_path, output_path):
    doc = docx.Document(docx_path)
    lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        style_name = para.style.name if para.style else ''
        
        if style_name.startswith('Heading 1'):
            lines.append(f'# {text}')
        elif style_name.startswith('Heading 2'):
            lines.append(f'## {text}')
        elif style_name.startswith('Heading 3'):
            lines.append(f'### {text}')
        elif style_name.startswith('Heading 4'):
            lines.append(f'#### {text}')
        else:
            lines.append(text)
    
    for table in doc.tables:
        lines.append('')
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if i == 0:
                lines.append('| ' + ' | '.join(cells) + ' |')
                lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
            else:
                lines.append('| ' + ' | '.join(cells) + ' |')
        lines.append('')
    
    markdown_content = '\n'.join(lines)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown_content)
    
    print(f"Successfully parsed docx to markdown: {output_path}")

if __name__ == '__main__':
    docx_path = r'E:\06 - 基于知识图谱的个性化荐书系统 - 接口文档.docx'
    output_path = r'd:\code\mybookhome\api-docs.md'
    parse_docx_to_markdown(docx_path, output_path)